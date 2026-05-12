"""
Tests for performance-review-generator task.
Verifies split placeholders, headers/footers, nested tables, and conditionals.
"""

import json
import re
from pathlib import Path

import pytest
from docx import Document

OUTPUT_FILE = "/root/performance_review_filled.docx"
DATA_FILE = "/root/review_data.json"


@pytest.fixture(scope="module")
def output_doc():
    """Load the output document (implicitly tests existence and validity)."""
    path = Path(OUTPUT_FILE)
    assert path.exists(), f"Output file not found: {OUTPUT_FILE}"
    return Document(OUTPUT_FILE)


@pytest.fixture(scope="module")
def review_data():
    """Load the review data."""
    with open(DATA_FILE) as f:
        return json.load(f)


def get_all_text(doc):
    """Extract all text from document including tables, headers, footers."""
    text_parts = []

    # Main paragraphs
    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Tables (including nested)
    def extract_from_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text)
                for nested in cell.tables:
                    extract_from_table(nested)

    for table in doc.tables:
        extract_from_table(table)

    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            text_parts.append(para.text)
        for para in section.footer.paragraphs:
            text_parts.append(para.text)

    return "\n".join(text_parts)


def get_nested_table_text(doc):
    """Get text from nested tables only."""
    text_parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for para in ncell.paragraphs:
                                text_parts.append(para.text)
    return "\n".join(text_parts)


# ============ No Remaining Placeholders ============


def test_no_remaining_placeholders(output_doc):
    """No {{...}} placeholders should remain anywhere in the document."""
    all_text = get_all_text(output_doc)
    matches = re.findall(r"\{\{[A-Z_]+\}\}", all_text)
    assert not matches, f"Unreplaced placeholders: {matches}"


# ============ Split Placeholder Tests ============

# Placeholders split across XML runs in the template
SPLIT_PLACEHOLDER_FIELDS = [
    "REVIEW_DATE",
    "EMPLOYEE_FULL_NAME",
    "EMPLOYEE_ID",
    "JOB_TITLE",
    "REVIEW_PERIOD",
    "OVERALL_RATING",
    "REVIEWER_NAME",
    "NEXT_REVIEW_DATE",
    "GOALS_MET",
    "SALARY_INCREASE",
]


@pytest.mark.parametrize("field", SPLIT_PLACEHOLDER_FIELDS)
def test_split_placeholder_replaced(output_doc, review_data, field):
    """Split placeholders should be correctly replaced."""
    all_text = get_all_text(output_doc)
    expected = review_data[field]
    assert expected in all_text, f"{field}='{expected}' not found in document"


# ============ Nested Table Tests ============

# Fields that appear in nested tables
NESTED_TABLE_FIELDS = [
    "JOB_TITLE",
    "DEPARTMENT",
    "OVERALL_RATING",
    "GOALS_MET",
    "ATTENDANCE_SCORE",
    "TEAMWORK_RATING",
    "TECHNICAL_RATING",
    "SALARY_INCREASE",
]


@pytest.mark.parametrize("field", NESTED_TABLE_FIELDS)
def test_nested_table_value(output_doc, review_data, field):
    """Values in nested tables should be correctly replaced."""
    nested_text = get_nested_table_text(output_doc)
    expected = review_data[field]
    assert expected in nested_text, f"{field}='{expected}' not in nested table"


# ============ Conditional Section Tests ============


def test_conditional_section(output_doc, review_data):
    """Conditional section should be handled (markers removed, content kept)."""
    all_text = get_all_text(output_doc)

    # IF markers should be removed
    assert "{{IF_BONUS}}" not in all_text, "{{IF_BONUS}} marker not removed"
    assert "{{END_IF_BONUS}}" not in all_text, "{{END_IF_BONUS}} marker not removed"

    # Content should be present (BONUS_ELIGIBLE is Yes)
    assert review_data["BONUS_AMOUNT"] in all_text, f"BONUS_AMOUNT='{review_data['BONUS_AMOUNT']}' not found"
    assert review_data["BONUS_PAYOUT_DATE"] in all_text, f"BONUS_PAYOUT_DATE='{review_data['BONUS_PAYOUT_DATE']}' not found"
